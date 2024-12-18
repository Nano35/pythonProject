import sqlite3
from email.policy import default
import requests
import os
import tempfile
from docx import Document
from dotenv import load_dotenv
from telebot import types
from docx.shared import Inches
load_dotenv()

import requests

import telebot
from telebot import types

import config

bot = telebot.TeleBot(config.get_api_key())
conn = sqlite3.connect('db/python.db', check_same_thread=False)
cursor = conn.cursor()

global fio
global birthday
global when_i_birth
global my_present
global my_future
global file_path

STATE_WAITING_FOR_ANSWER = 0
ANSWER_RECEIVED = 1

@bot.message_handler(commands=['start']) #команды
def send_welcome(message):
    bot.reply_to(message, """\
Привет! Я бот фонда: "Измени Одну Жизнь"
Я здесь чтобы помочь приемным родителям собрать всю информацию о прошлом их приемного ребенка и его настоящем уже в новой семье\
""")
    inline_markup = types.InlineKeyboardMarkup()
    keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
    btn1 = types.InlineKeyboardButton(text='Наш сайт', url='https://changeonelife.ru/about/mission/')
    hi_btn = types.KeyboardButton("Привет")
    inline_markup.add(btn1)
    keyboard.add(hi_btn)
    bot.send_message(message.from_user.id, "По кнопке ниже можно перейти на наш сайт", reply_markup = inline_markup)
    bot.send_message(message.from_user.id, "Выберите действия, которые вы хотите совершить:" , reply_markup=keyboard)

@bot.message_handler(content_types=['text'])
def get_text_messages(message):
   if message.text == 'Привет':
       if check_user(message.from_user.id):
           bot.send_message(message.from_user.id, 'Привет! Вы уже были зарегистрированы!')
       else:
        add_user(message.from_user.id, message.from_user.first_name, message.from_user.last_name, message.from_user.username)
        bot.send_message(message.from_user.id, 'Привет! Ваше имя добавлено в базу данных!')
       keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
       reg_btn = types.KeyboardButton("📕 Собрать книгу жизни")
       keyboard.add(reg_btn)
       bot.send_message(message.from_user.id, "Выберите действие:", reply_markup=keyboard)
   elif message.text == '📕 Собрать книгу жизни':
       bot.send_message(message.from_user.id, "Напишите ФИО ребенка")
       bot.register_next_step_handler(message, add_child_data)


def add_child_data(message):
    global fio
    fio = message.text
    bot.reply_to(message, f"ФИО ребенка: {fio}")
    bot.send_message(message.from_user.id, "Загрузите фото ребенка")
    bot.register_next_step_handler(message, add_child_photo)

def add_child_photo(message):
    file = bot.get_file(message.photo[-1].file_id)
    global file_path
    file_path = file.file_path  # Путь к сохраненному файлу который сохраняется в бд
    global file_url
    file_url = f"https://api.telegram.org/file/bot{config.get_api_key()}/{file_path}" # Реальный путь (можете по нему перейти там в консоли он выводится) НЕ ИСПОЛЬЗОВАТЬ В БД ТАК КАК ПАЛИТСЯ BOT API KEY
    print(file_url)
    response = requests.get(file_url, stream=True)
    response.raise_for_status()  # Проверяем на ошибки

    with open(f"photo_{file.file_id}.jpg", "wb") as f:  # Сохраняем файл на сервер
        for chunk in response.iter_content(chunk_size=8192):
            f.write(chunk)
    bot.send_message(message.from_user.id, "Напишите день рождения ребенка:")
    bot.register_next_step_handler(message, add_birthday)

def add_birthday(message):
    global birthday
    birthday = message.text
    bot.reply_to(message, f"День рождения ребенка: {birthday}")
    bot.send_message(message.from_user.id, """*Моё прошлое*
Расскажите о прошлом ребенка \(о том когда он родился и где он родился\)""", parse_mode="MarkdownV2")
    bot.send_message(message.from_user.id, "Напишите информацию о том когда он родился и где он родился", parse_mode="MarkdownV2")
    bot.register_next_step_handler(message, add_when_i_birth)

def add_when_i_birth(message):
    global when_i_birth
    when_i_birth = message.text
    bot.send_message(message.from_user.id, """*Моё настоящее*
Расскажите о настоящем ребенка \(о том где он сейчас живет, люди которые с ним, распорядок дня и его свободное время\)""", parse_mode="MarkdownV2")
    bot.register_next_step_handler(message, add_my_present)

def add_my_present(message):
    global my_present
    my_present = message.text
    bot.send_message(message.from_user.id, """*Моё будущее*
Расскажите о будущем ребенка \(о том кем он хочет быть, о его будущей семье, о его будущем доме\)""", parse_mode="MarkdownV2")
    bot.register_next_step_handler(message, add_child_db)


#СОЗДАНИЕ ДОКУМЕНТА
def finish(message):
    chat_id = message.chat.id
    keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
    button = types.KeyboardButton(text="Создать документ word")
    keyboard.add(button)

    msg = bot.send_message(
        chat_id,
        "Нажмите на кнопку, чтобы создать документ",
        reply_markup=keyboard
        )
    bot.register_next_step_handler(msg, intro_text_handler)
#СОЗДАНИЕ ДОКУМЕНТА
def intro_text_handler(message):
    """Обработчик текста"""
    chat_id = message.chat.id
    # Генерируем документ и отправляем его пользователю
    try:
        document = generate_document()
        msg = bot.send_document(chat_id, document)
        bot.register_next_step_handler(msg, intro_text_handler)
    except Exception as e:
        bot.send_message(chat_id, "Ошибка генерации документа: " + str(e))

#ДЕЛАЕТ ИЗОБРАЖЕНИЕ ИЗ ССЫЛКИ
def insert_image_from_url(document, image_url, width=None):
    """Вставляет изображение в документ docx по URL-ссылке.

    Args:
        document: Объект Document из библиотеки python-docx.
        image_url: URL-ссылка на изображение.
        width: Ширина изображения в дюймах (опционально).
    """

    try:
        response = requests.get(image_url, stream=True)
        response.raise_for_status()

        with open('temp_image.jpg', 'wb') as out_file:
            for chunk in response.iter_content(chunk_size=8192):
                out_file.write(chunk)

        document.add_picture('temp_image.jpg', width=Inches(width) if width else None)
        os.remove('temp_image.jpg')

    except requests.exceptions.RequestException as e:
        print(f"Ошибка при загрузке изображения: {e}")
    except Exception as e:
        print(f"Ошибка при вставке изображения: {e}")
#СОЗДАНИЕ ДОКУМЕНТА
def generate_document():
    """Генератор Word документа"""
    try:
        doc = Document()
        image_url = file_url
        doc.add_heading('Книга жизни', level=0)
        doc.add_paragraph('Это инструмент, который позволяет приемным родителям собрать всю информацию о прошлом их приемного ребенка и его настоящем уже в новой семье и оформить это в виде "Книги жизни" - альбома с текстами и фотографиями')
        doc.add_heading('Вводная часть:', level=1)
        doc.add_paragraph(fio, style = 'List Number')
        doc.add_paragraph(birthday, style = 'List Number')
        insert_image_from_url(doc, image_url, width=1.25)
        doc.add_heading('Мое прошлое:', level=1)
        doc.add_paragraph(when_i_birth)
        doc.add_heading('Мое настоящее:', level=1)
        doc.add_paragraph(my_present)
        doc.add_heading('Мое будущее:', level=1)
        doc.add_paragraph(my_future)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            document_name = f.name
            doc.save(document_name)
            return open(document_name, 'rb')
    except Exception as e:
        raise Exception("Ошибка генерации документа: " + str(e))




def add_child_db(message): # Сохраняем path в базу данных
    global my_future
    my_future = message.text
    resume = when_i_birth + "\n" + my_present + "\n" + my_future

    cursor.execute("INSERT INTO child (fio, birthday, resume, photo) VALUES (?, ?, ?, ?)",
                   (fio, birthday, resume, file_path))  # user_id - ID пользователя из Telegram
    conn.commit()
    conn.close()
    #отсылка на функцию ворда
    bot.register_next_step_handler(message, finish)
def check_user(user_id):
    cursor.execute('SELECT 1 FROM test WHERE user_id = ? LIMIT 1', (user_id,))
    result = cursor.fetchone()
    return result

def add_user(user_id: int, user_name: str, user_surname: str, username: str):

    cursor.execute('INSERT INTO test (user_id, user_name, user_surname, username) VALUES (?, ?, ?, ?)',
                   (user_id, user_name, user_surname, username))
    conn.commit()













if __name__ == '__main__':
    bot.polling()